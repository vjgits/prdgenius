"""
PRDGenius – production FastAPI back-end
"""
import io, os, re, uuid, json, hashlib, sqlite3, logging, asyncio
from datetime import datetime, timedelta
from contextlib import asynccontextmanager
from typing import Optional

from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse, PlainTextResponse, Response
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
import anthropic
import stripe
from docx import Document
from docx.shared import Inches, Pt

# ─── Logging ─────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Rate limiting store ──────────────────────────────────────────────────────
_rate_store: dict = {}

def check_rate_limit(key: str, max_calls: int, window_seconds: int) -> bool:
    now = datetime.utcnow().timestamp()
    calls = _rate_store.get(key, [])
    calls = [t for t in calls if now - t < window_seconds]
    if len(calls) >= max_calls:
        return False
    calls.append(now)
    _rate_store[key] = calls
    return True

# ─── Allowed email domains (blocks throwaway email abuse) ─────────────────────
ALLOWED_EMAIL_DOMAINS = {
    "gmail.com", "yahoo.com", "outlook.com", "hotmail.com",
    "icloud.com", "protonmail.com", "me.com", "mac.com",
    "live.com", "msn.com", "ymail.com", "googlemail.com",
    "aol.com", "proton.me", "fastmail.com", "hey.com",
    "zoho.com", "mail.com", "pm.me", "tutanota.com",
    "prdgenius.ai"
}

# ─── Config ───────────────────────────────────────────────────────────────────
DB_PATH                = os.getenv("DB_PATH", "/data/prd_genius.db")
PASSWORD_SALT          = os.getenv("PASSWORD_SALT", "prdgenius_s3cur3_s4lt_2024")
PRODUCTION             = os.getenv("PRODUCTION", "false").lower() == "true"
# Credit system — Brief=3 / Medium=4 / Extensive=6 per PRD
CREDIT_COST   = {"brief": 3, "medium": 4, "extensive": 6}
FREE_CREDITS  = 6    # enough for 1 PRD of any size (enforced by FREE_PRD_LIMIT below)
FREE_PRD_LIMIT = 1   # free users get exactly 1 PRD total
FREE_FORMATS = ["microsoft", "tesla", "linear"]  # formats available on free plan
PRO_CREDITS   = 120  # 40 Brief / 30 Medium / 20 Extensive
stripe.api_key         = os.getenv("STRIPE_SECRET_KEY", "")
STRIPE_PUBLISHABLE_KEY = os.getenv("STRIPE_PUBLISHABLE_KEY", "")
STRIPE_WEBHOOK_SECRET  = os.getenv("STRIPE_WEBHOOK_SECRET", "")
STRIPE_PRICE_ID        = os.getenv("STRIPE_PRICE_ID", "")
STRIPE_YEARLY_PRICE_ID = os.getenv("STRIPE_YEARLY_PRICE_ID", "")
BASE_URL               = os.getenv("BASE_URL", "http://127.0.0.1:8000")
ANTHROPIC_API_KEY      = os.getenv("ANTHROPIC_API_KEY", "")

# ─── Database ─────────────────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    return conn

def init_db():
    conn = get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id                     TEXT PRIMARY KEY,
            name                   TEXT NOT NULL,
            email                  TEXT UNIQUE NOT NULL,
            password_hash          TEXT NOT NULL,
            plan                   TEXT NOT NULL DEFAULT 'free',
            role                   TEXT NOT NULL DEFAULT 'user',
            stripe_customer_id     TEXT,
            stripe_subscription_id TEXT,
            prds_used_this_month   INTEGER NOT NULL DEFAULT 0,
            credits_used           INTEGER NOT NULL DEFAULT 0,
            credits_month          TEXT    NOT NULL DEFAULT '',
            created_at             TEXT NOT NULL DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS prds (
            id           TEXT PRIMARY KEY,
            user_id      TEXT NOT NULL,
            title        TEXT NOT NULL,
            content      TEXT NOT NULL,
            format_style TEXT NOT NULL,
            created_at   TEXT NOT NULL DEFAULT (datetime('now'))
        );
        CREATE TABLE IF NOT EXISTS sessions (
            token      TEXT PRIMARY KEY,
            user_id    TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            created_at TEXT NOT NULL DEFAULT (datetime('now'))
        );
    """)
    conn.commit()
    existing = [r[1] for r in conn.execute("PRAGMA table_info(prds)").fetchall()]
    for col in ["target_users","key_features","success_metrics","company_stage","additional_context"]:
        if col not in existing:
            conn.execute(f"ALTER TABLE prds ADD COLUMN {col} TEXT DEFAULT ''")
    # Migrate users table — add credit columns for existing installs
    user_cols = [r[1] for r in conn.execute("PRAGMA table_info(users)").fetchall()]
    for col, defn in [
        ("credits_used",  "INTEGER NOT NULL DEFAULT 0"),
        ("credits_month", "TEXT NOT NULL DEFAULT ''"),
    ]:
        if col not in user_cols:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col} {defn}")
    conn.commit()
    admin_email = "admin@prdgenius.ai"
    admin_pw    = hash_password("Chocolate47##")
    existing    = conn.execute("SELECT id FROM users WHERE email=?", (admin_email,)).fetchone()
    if not existing:
        conn.execute(
            "INSERT INTO users (id, name, email, password_hash, plan, role) VALUES (?,?,?,?,?,?)",
            (str(uuid.uuid4()), "Admin", admin_email, admin_pw, "admin", "admin")
        )
    else:
        conn.execute("UPDATE users SET role='admin', plan='admin' WHERE email=?", (admin_email,))
    conn.commit()
    conn.close()

# ─── Auth helpers ─────────────────────────────────────────────────────────────
def hash_password(password: str) -> str:
    return hashlib.sha256(f"{PASSWORD_SALT}{password}".encode()).hexdigest()

def create_session(user_id: str) -> str:
    token   = str(uuid.uuid4())
    expires = (datetime.utcnow() + timedelta(days=30)).isoformat()
    conn    = get_db()
    conn.execute("INSERT INTO sessions (token, user_id, expires_at) VALUES (?,?,?)",
                 (token, user_id, expires))
    conn.commit(); conn.close()
    return token

def get_current_user(request: Request) -> Optional[dict]:
    token = request.cookies.get("session")
    if not token: return None
    conn = get_db()
    row  = conn.execute(
        "SELECT s.user_id, s.expires_at FROM sessions s WHERE s.token=?", (token,)
    ).fetchone()
    if not row: conn.close(); return None
    if datetime.fromisoformat(row["expires_at"]) < datetime.utcnow():
        conn.execute("DELETE FROM sessions WHERE token=?", (token,))
        conn.commit(); conn.close(); return None
    user = conn.execute("SELECT * FROM users WHERE id=?", (row["user_id"],)).fetchone()
    conn.close()
    return dict(user) if user else None

# ─── Plan helpers ─────────────────────────────────────────────────────────────
def get_credit_limit(user: dict) -> int:
    if user.get("role") == "admin" or user.get("plan") in ("pro", "admin", "yearly"):
        return PRO_CREDITS
    return FREE_CREDITS

def reset_credits_if_new_month(user: dict, conn) -> dict:
    """Reset credits_used if we're in a new calendar month. Returns refreshed user dict."""
    current_month = datetime.utcnow().strftime("%Y-%m")
    if user.get("credits_month", "") != current_month:
        conn.execute(
            "UPDATE users SET credits_used=0, prds_used_this_month=0, credits_month=? WHERE id=?",
            (current_month, user["id"])
        )
        conn.commit()
        row = conn.execute("SELECT * FROM users WHERE id=?", (user["id"],)).fetchone()
        return dict(row) if row else user
    return user

# ─── Security Middleware ──────────────────────────────────────────────────────
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"]  = "nosniff"
        response.headers["X-Frame-Options"]          = "DENY"
        response.headers["X-XSS-Protection"]         = "1; mode=block"
        response.headers["Referrer-Policy"]           = "strict-origin-when-cross-origin"
        return response

# ─── DOCX helpers ─────────────────────────────────────────────────────────────
def _add_inline_formatting(paragraph, text: str):
    for part in re.split(r'(\*\*.*?\*\*|`.*?`)', text):
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'; run.font.size = Pt(9)
        else:
            paragraph.add_run(part)

def markdown_to_docx(md: str) -> bytes:
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1.2)
    for s in md.split('\n'):
        if   s.startswith('# '):   doc.add_heading(s[2:],  level=1)
        elif s.startswith('## '):  doc.add_heading(s[3:],  level=2)
        elif s.startswith('### '): doc.add_heading(s[4:],  level=3)
        elif s.startswith(('- ','* ')):
            _add_inline_formatting(doc.add_paragraph(style='List Bullet'), s[2:].strip())
        elif re.match(r'^\d+\. ', s):
            _add_inline_formatting(doc.add_paragraph(style='List Number'), re.sub(r'^\d+\. ','',s))
        elif s.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
            doc.add_paragraph().add_run(s[2:]).italic = True
        elif s.strip() == '': doc.add_paragraph()
        else: _add_inline_formatting(doc.add_paragraph(), s)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

# ─── App ─────────────────────────────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app):
    init_db()
    yield

app = FastAPI(title="PRDGenius", lifespan=lifespan, docs_url=None, redoc_url=None)
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True,
                   allow_methods=["GET","POST","DELETE"], allow_headers=["*"])
templates = Jinja2Templates(directory="templates")

# ─── SEO ──────────────────────────────────────────────────────────────────────
SITE_URL = "https://prdgenius.up.railway.app"

@app.get("/robots.txt", response_class=PlainTextResponse)
async def robots():
    return "\n".join([
        "User-agent: *",
        "Allow: /",
        "Disallow: /api/",
        "Disallow: /app",
        "Disallow: /prd/",
        f"Sitemap: {SITE_URL}/sitemap.xml",
    ])

@app.get("/sitemap.xml")
async def sitemap():
    urls = [
        {"loc": SITE_URL + "/",        "priority": "1.0", "changefreq": "weekly"},
        {"loc": SITE_URL + "/signup",  "priority": "0.9", "changefreq": "monthly"},
        {"loc": SITE_URL + "/login",   "priority": "0.5", "changefreq": "monthly"},
    ]
    lines = ['<?xml version="1.0" encoding="UTF-8"?>',
             '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">']
    for u in urls:
        lines += [
            "  <url>",
            f"    <loc>{u['loc']}</loc>",
            f"    <priority>{u['priority']}</priority>",
            f"    <changefreq>{u['changefreq']}</changefreq>",
            "  </url>",
        ]
    lines.append("</urlset>")
    return Response("\n".join(lines), media_type="application/xml")

# ─── Pages ────────────────────────────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def landing(request: Request):
    user = get_current_user(request)
    if user: return RedirectResponse("/app")
    return templates.TemplateResponse("landing.html", {"request": request})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    if get_current_user(request): return RedirectResponse("/app")
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/signup", response_class=HTMLResponse)
async def signup_page(request: Request):
    if get_current_user(request): return RedirectResponse("/app")
    return templates.TemplateResponse("signup.html", {"request": request})

@app.get("/app", response_class=HTMLResponse)
async def app_page(request: Request):
    user = get_current_user(request)
    if not user: return RedirectResponse("/login")
    conn = get_db()
    user = reset_credits_if_new_month(user, conn)
    history_limit = 40 if user.get("plan") in ("pro", "admin", "yearly") else 5
    prds = [dict(p) for p in conn.execute(
        "SELECT * FROM prds WHERE user_id=? ORDER BY created_at DESC LIMIT ?", (user["id"], history_limit)
    ).fetchall()]
    conn.close()
    credit_limit    = get_credit_limit(user)
    credits_used    = user.get("credits_used", 0)
    credits_remaining = max(0, credit_limit - credits_used)
    is_free = user.get("plan", "free") not in ("pro", "admin", "yearly")
    at_limit = (is_free and user.get("prds_used_this_month", 0) >= FREE_PRD_LIMIT) or \
               (not is_free and credits_remaining < 3)
    return templates.TemplateResponse("app.html", {
        "request": request, "user": user, "prds": prds,
        "last_prd": prds[0] if prds else None,
        "at_limit": at_limit,
        "credit_limit": credit_limit,
        "credits_used": credits_used,
        "credits_remaining": credits_remaining,
        "stripe_pub_key": STRIPE_PUBLISHABLE_KEY,
        "free_credits": FREE_CREDITS,
        "pro_credits": PRO_CREDITS,
        "free_formats": FREE_FORMATS,
    })

@app.get("/upgrade", response_class=HTMLResponse)
async def upgrade_page(request: Request):
    user = get_current_user(request)
    if not user: return RedirectResponse("/login")
    return templates.TemplateResponse("upgrade.html", {
        "request": request, "user": user, "stripe_pub_key": STRIPE_PUBLISHABLE_KEY,
        "free_credits": FREE_CREDITS, "pro_credits": PRO_CREDITS,
    })

@app.get("/upgrade/success", response_class=HTMLResponse)
async def upgrade_success(request: Request):
    user = get_current_user(request)
    if not user: return RedirectResponse("/login")
    return templates.TemplateResponse("upgrade_success.html", {"request": request, "user": user})

@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    user = get_current_user(request)
    if not user or user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    return templates.TemplateResponse("admin.html", {"request": request, "user": user})

# ─── Auth API ─────────────────────────────────────────────────────────────────
@app.post("/api/signup")
async def api_signup(
    request:  Request,
    name:     str = Form(...),
    email:    str = Form(...),
    password: str = Form(...)
):
    client_ip = request.client.host
    if not check_rate_limit(f"signup:{client_ip}", max_calls=3, window_seconds=600):
        return JSONResponse(
            {"error": "Too many signup attempts from your location. Please try again later."},
            status_code=429
        )
    email  = email.strip().lower()
    domain = email.split("@")[-1] if "@" in email else ""
    if domain not in ALLOWED_EMAIL_DOMAINS:
        return JSONResponse(
            {"error": f"Signups require a major email provider (Gmail, Yahoo, Outlook, iCloud, etc.). '{domain}' is not supported."},
            status_code=400
        )
    if len(password) < 8:
        return JSONResponse({"error": "Password must be at least 8 characters."}, status_code=400)
    conn     = get_db()
    existing = conn.execute("SELECT id FROM users WHERE email=?", (email,)).fetchone()
    if existing:
        conn.close()
        return JSONResponse({"error": "Email already registered."}, status_code=409)
    user_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO users (id, name, email, password_hash, plan, role) VALUES (?,?,?,?,?,?)",
        (user_id, name.strip(), email, hash_password(password), "free", "user")
    )
    conn.commit(); conn.close()
    token    = create_session(user_id)
    response = JSONResponse({"success": True, "redirect": "/app"})
    response.set_cookie("session", token, httponly=True, secure=PRODUCTION,
                        max_age=60*60*24*30, samesite="lax")
    return response

@app.post("/api/login")
async def api_login(
    request:  Request,
    email:    str = Form(...),
    password: str = Form(...)
):
    client_ip = request.client.host
    if not check_rate_limit(f"login:{client_ip}", max_calls=10, window_seconds=300):
        return JSONResponse(
            {"error": "Too many login attempts. Please try again in 5 minutes."},
            status_code=429
        )
    email = email.strip().lower()
    conn  = get_db()
    user  = conn.execute("SELECT * FROM users WHERE email=?", (email,)).fetchone()
    conn.close()
    if not user or user["password_hash"] != hash_password(password):
        return JSONResponse({"error": "Invalid email or password."}, status_code=401)
    token    = create_session(user["id"])
    response = JSONResponse({"success": True, "redirect": "/app"})
    response.set_cookie("session", token, httponly=True, secure=PRODUCTION,
                        max_age=60*60*24*30, samesite="lax")
    return response

@app.post("/api/logout")
async def api_logout(request: Request):
    token = request.cookies.get("session")
    if token:
        conn = get_db()
        conn.execute("DELETE FROM sessions WHERE token=?", (token,))
        conn.commit(); conn.close()
    response = RedirectResponse("/", status_code=302)
    response.delete_cookie("session")
    return response

# ─── PRD API ──────────────────────────────────────────────────────────────────
@app.post("/api/check-limit")
async def check_limit(request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    conn = get_db()
    user = reset_credits_if_new_month(user, conn)
    conn.close()
    credit_limit      = get_credit_limit(user)
    credits_used      = user.get("credits_used", 0)
    credits_remaining = max(0, credit_limit - credits_used)
    plan = user.get("plan", "free")
    name = "Pro" if plan in ("pro","yearly") else ("Admin" if plan == "admin" else "Free")
    return JSONResponse({
        "allowed": credits_remaining >= 3,
        "credits_used": credits_used,
        "credit_limit": credit_limit,
        "credits_remaining": credits_remaining,
        "plan": plan, "plan_name": name,
        "message": f"Monthly credit limit reached on {name} plan." if credits_remaining < 3 else ""
    })

@app.post("/api/generate")
async def generate_prd(request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    if not check_rate_limit(f"gen:{user['id']}", max_calls=5, window_seconds=60):
        return JSONResponse({"error": "Too many requests. Please wait a moment."}, status_code=429)

    data               = await request.json()
    product_name       = data.get("product_name",    "").strip()
    problem            = data.get("problem",          "").strip()
    target_users       = data.get("target_users",     "").strip()
    key_features       = data.get("key_features",     "").strip()
    success_metrics    = data.get("success_metrics",  "").strip()
    format_style       = data.get("format_style",     "google").strip()
    company_stage      = data.get("company_stage",    "").strip()
    additional_context = data.get("additional_context", data.get("context", "")).strip()

    if not product_name or not problem:
        return JSONResponse({"error": "Product name and problem statement are required."}, status_code=400)

    prd_size = data.get("prd_size", "medium").strip()
    if prd_size not in CREDIT_COST:
        prd_size = "medium"
    credit_cost = CREDIT_COST[prd_size]

    # Monthly reset + credit check
    conn = get_db()
    user = reset_credits_if_new_month(user, conn)
    conn.close()
    credit_limit      = get_credit_limit(user)
    credits_used      = user.get("credits_used", 0)
    credits_remaining = credit_limit - credits_used
    is_free = user.get("plan", "free") not in ("pro", "admin", "yearly")
    free_limit_hit = is_free and user.get("prds_used_this_month", 0) >= FREE_PRD_LIMIT
    if free_limit_hit or credits_remaining < credit_cost:
        return JSONResponse(
            {"error": "You've reached your PRD limit for this month.", "upgrade": is_free},
            status_code=403
        )

    # Free users restricted to 3 formats
    if is_free and format_style not in FREE_FORMATS:
        return JSONResponse(
            {"error": "This PRD format is only available on the Pro plan. Upgrade to access all 24 formats.", "upgrade": True},
            status_code=403
        )

    # ── Format instructions for each style ────────────────────────────────────
    format_instructions = {
        "google": (
            "Write a comprehensive Google-style PRD. Be data-driven, precise, and cross-functionally "
            "aligned. Cover goals, personas, requirements, success metrics, technical considerations, "
            "launch plan, and risk analysis with the depth expected of a senior PM."
        ),
        "amazon": (
            "Write an Amazon Working Backwards PRD. Begin with an internal PRESS RELEASE (2 crisp "
            "paragraphs: customer, problem, solution, key benefit). Follow with an internal FAQ "
            "(customer questions first, then business/technical). Then write the full requirements. "
            "Every decision must trace back to the customer outcome."
        ),
        "linear": (
            "Write a Linear-style agile PRD. Be ruthlessly concise and sprint-ready. Lead with a "
            "sharp problem statement, proposed solution, and explicit acceptance criteria per "
            "requirement. Engineers must be able to start building from this document immediately. "
            "Cut all fluff."
        ),
        "jtbd": (
            "Write a Jobs-to-be-Done (JTBD) PRD. Anchor every requirement to the specific job the "
            "user is trying to accomplish. Use the structure: 'When [situation], I want to "
            "[motivation], so I can [outcome].' Define success entirely in terms of user outcomes, "
            "not feature outputs. Popularised by Clayton Christensen."
        ),
        "hypothesis": (
            "Write a Lean Hypothesis-driven PRD. Frame every feature as a testable bet: 'We believe "
            "[feature] will [outcome] for [user segment]. We will know this is true when [measurable "
            "signal].' Make every assumption explicit. Define what would prove or disprove each one. "
            "Ideal for early-stage or experimental features where you are still learning."
        ),
        "rfc": (
            "Write a Technical RFC (Request for Comments) PRD. Lead with technical context, system "
            "design options, and engineering tradeoffs. Include API contracts, data models, "
            "performance requirements, security considerations, and rollback plan. Written for the "
            "engineering team building it — not the stakeholders approving it."
        ),
        # ── Company styles ────────────────────────────────────────────────────
        "apple": (
            "Write an Apple-style PRD. Design and simplicity are the highest-order requirements — "
            "if the user experience needs a manual, you have failed. Describe experiences, not feature "
            "lists. Every requirement must answer: 'Is this magical? Is this the simplest it can be?' "
            "Apply Steve Jobs's 'say no to a thousand things' principle: cut ruthlessly. Hardware-"
            "software integration, privacy by default, and ecosystem coherence are non-negotiable."
        ),
        "microsoft": (
            "Write a Microsoft-style PRD. Enterprise-grade reliability, backward compatibility, and "
            "accessibility (WCAG 2.1 AA minimum) are first-class requirements. Address IT admin "
            "control, security and compliance (SOC2, ISO 27001), and integration with the Microsoft "
            "365 / Azure / Teams ecosystem. Quantify success with enterprise adoption metrics, NPS "
            "from both end users and IT admins, and power-user efficiency gains."
        ),
        "meta": (
            "Write a Meta-style PRD. Every requirement must leverage the social graph and drive "
            "network effects. Frame success metrics around DAU, MAU, engagement rate, and social "
            "sharing velocity — not just usage. A/B testing plan is mandatory: nothing ships without "
            "a statistically valid experiment design. Optimise for time-to-first-social-interaction "
            "and viral coefficient. Address News Feed / algorithmic distribution implications."
        ),
        "nvidia": (
            "Write an NVIDIA-style PRD. Lead with hard performance specifications: throughput "
            "(tokens/s, TFLOPS), latency (P50/P99), memory bandwidth, and power envelope. Frame "
            "requirements around the CUDA and developer toolchain ecosystem. Address hardware-software "
            "co-design constraints and include benchmark comparisons against competitive solutions. "
            "Developer experience and SDK ergonomics are as important as raw performance."
        ),
        "openai": (
            "Write an OpenAI-style PRD. Pair every capability requirement with a corresponding safety "
            "and evaluation requirement — they carry equal weight. Include red-teaming criteria, "
            "misuse vector analysis, and model evaluation benchmarks. Address responsible deployment "
            "milestones and staged rollout gates. Frame success as capability × safety, not capability "
            "alone. Define what 'ready to ship' means from both a product and safety standpoint."
        ),
        "anthropic": (
            "Write an Anthropic-style PRD grounded in Constitutional AI principles. The safety case "
            "must be as developed as the product case — this is non-negotiable. Every feature must "
            "include a harmlessness assessment, an honesty evaluation, and a helpfulness measure. "
            "Address responsible scaling policy implications and include evaluation criteria for both "
            "capability and alignment. Pair ambition with epistemic humility: be explicit about what "
            "you do not yet know."
        ),
        "tesla": (
            "Write a Tesla-style PRD using Elon Musk's 5-step engineering process: (1) Challenge "
            "every requirement — if you can't question it, you can't delete it. (2) Delete any part, "
            "process, or spec that can be removed — you can always add back. (3) Simplify and "
            "optimise only after deleting. (4) Accelerate cycle time. (5) Automate. Requirements must "
            "be OTA-deliverable, vertically integrated, and justified by physics, not convention. "
            "Include a COGS target alongside the feature requirements."
        ),
        "mercedes": (
            "Write a Mercedes-Benz-style PRD. Premium quality and German engineering precision define "
            "every requirement. Safety standards (ISO 26262 ASIL rating for automotive, GDPR for "
            "data) are non-negotiable. User experience must reflect the luxury segment: every "
            "interaction crafted, not engineered. Include quality validation gates, homologation "
            "requirements, and long-term reliability KPIs alongside functional requirements. "
            "Durability and brand coherence carry as much weight as new capabilities."
        ),
        "honda": (
            "Write a Honda-style PRD using The Honda Way. Apply kaizen (continuous improvement) "
            "thinking: frame requirements as disciplined improvements to a proven, reliable baseline. "
            "Manufacturing feasibility and total cost of ownership are first-class inputs — involve "
            "manufacturing in the design phase, not after. Include quality standards (defect rate "
            "targets), serviceability requirements, and global regulatory compliance across key "
            "markets. Reliability over novelty."
        ),
        "elililly": (
            "Write an Eli Lilly-style PRD for a regulated healthcare product. All requirements must "
            "align with FDA 21 CFR and ICH E6 guidelines. Structure around clinical evidence: define "
            "primary and secondary endpoints, safety profile thresholds, and pharmacovigilance "
            "obligations. Include health economics framing (QALY, cost-effectiveness). Patient "
            "outcomes are the north star — every feature must trace back to a measurable impact on "
            "patient health or care team efficiency. Regulatory strategy is a requirement, not an "
            "afterthought."
        ),
        "novartis": (
            "Write a Novartis-style PRD grounded in patient-centricity and clinical evidence. Lead "
            "with the unmet medical need and the clinical development rationale. Address regulatory "
            "strategy across FDA, EMA, and PMDA pathways. Include pharmacovigilance obligations, "
            "medical affairs input, and real-world evidence strategy. Frame success metrics around "
            "clinical outcomes, patient quality of life (PRO instruments), and market access "
            "criteria. Every product decision must withstand scientific and regulatory scrutiny."
        ),
        "exxon": (
            "Write an ExxonMobil-style PRD for an energy sector product. Operational safety (Process "
            "Safety Management per OSHA 1910.119) is the highest priority requirement — no exceptions. "
            "Include a comprehensive HSE (Health, Safety, Environment) impact assessment. Address "
            "regulatory compliance (EPA, FERC), CAPEX/OPEX analysis, and upstream/downstream "
            "integration implications. Frame success in terms of operational efficiency gains, "
            "emissions reduction, and long-term asset value. ESG reporting alignment is mandatory."
        ),
        "chevron": (
            "Write a Chevron-style PRD. Operational safety and Process Safety Management (PSM) "
            "standards take absolute precedence. Address environmental stewardship requirements, "
            "energy transition alignment (net-zero roadmap), and ESG impact measurement. Include "
            "CAPEX/OPEX analysis, regulatory compliance (EPA, state regulators), and supply chain "
            "resilience considerations. Success metrics must include safety KPIs (TRIR, process "
            "safety events), environmental performance indicators, and financial returns."
        ),
        "maersk": (
            "Write a MAERSK-style PRD for a global logistics and supply chain product. Focus on "
            "supply chain resilience, end-to-end visibility, and decarbonisation targets (net-zero "
            "2040 commitment). Requirements must address multi-modal transport integration (ocean, "
            "inland, air), real-time tracking and exception management, and trade lane economics. "
            "Digital transformation of physical logistics is the core theme. Include sustainability "
            "metrics (CO2 per TEU-km) alongside operational KPIs."
        ),
        "jpmorgan": (
            "Write a JPMorgan Chase-style PRD. Regulatory compliance (Basel III, Dodd-Frank, MiFID "
            "II, GDPR) is a first-class product requirement, not a legal afterthought. Every feature "
            "must include a risk assessment, full audit trail specification, financial controls "
            "design, and model risk management consideration where applicable. Frame success metrics "
            "around risk-adjusted returns, compliance adherence rates, and operational resilience "
            "targets. Security and fraud prevention are requirements, not features."
        ),
        "visa": (
            "Write a Visa-style PRD for a global payments product. Five-nines (99.999%) availability "
            "and sub-100ms authorisation latency are baseline requirements, not stretch goals. "
            "Address fraud detection and prevention at global scale, PCI-DSS and PSD2 compliance, "
            "and tokenisation architecture. Frame success metrics around authorisation rate, fraud "
            "basis points, and merchant/cardholder acceptance. Global regulatory variation (EMV, "
            "local payment schemes) must be addressed per region."
        ),
        "spacex": (
            "Write a SpaceX-style PRD using first-principles thinking. Question every physical and "
            "cost constraint from the ground up — industry convention is not a valid requirement "
            "source. Justify requirements by physics and economics, not precedent. Include aggressive "
            "cost-per-unit targets, full or partial reusability requirements, and iteration cycle "
            "time goals. 'Faster than expected' is a success criterion. Address reliability and "
            "safety with FMEA discipline, but never let safety theatre slow rational progress."
        ),
        "boeing": (
            "Write a Boeing-style PRD for an aerospace or defence product. Safety-critical "
            "requirements are absolute and must reference FAA/EASA airworthiness standards. Assign "
            "DO-178C software criticality levels (DAL A–E) for all software components. Include "
            "system redundancy architecture, FMEA (Failure Mode and Effects Analysis) summary, and "
            "design assurance requirements. Certification timeline is a product requirement. "
            "Traceability from requirement to test case is mandatory."
        ),
        # legacy aliases kept for backwards compat
        "standard":  "Write a comprehensive, well-structured PRD with clear sections.",
        "lean":      "Write a concise lean PRD focusing on essentials only.",
        "agile":     "Write an agile-style PRD with user stories and acceptance criteria.",
        "technical": "Write a technical PRD with system requirements and engineering details.",
    }.get(format_style, "Write a comprehensive, well-structured PRD.")

    # Per-section word cap guarantees all 10 sections + TOC fit within token budget
    size_config = {
        "brief":     (2000, 120),
        "medium":    (4000, 250),
        "extensive": (6000, 400),
    }
    max_tok, words_per_section = size_config[prd_size]

    # Each format gets unique section headings — this is what makes outputs genuinely different
    format_sections_map = {
        "google": """1. Executive Summary\n2. Problem Statement & Background\n3. Goals & Key Results (OKRs)\n4. User Personas & Target Audience\n5. Feature Requirements (P0 / P1 / P2)\n6. User Stories\n7. Technical Considerations\n8. Launch Plan & Timeline\n9. Risks & Mitigations\n10. Open Questions""",
        "amazon": """1. Press Release (customer-facing announcement)\n2. Customer FAQ\n3. Internal / Business FAQ\n4. Customer Personas\n5. Working Backwards Requirements\n6. Feature Specification\n7. Success Metrics (customer outcomes)\n8. Implementation & Dependencies\n9. Risks & Open Questions\n10. Working Backwards Checklist""",
        "linear": """1. Problem\n2. Why Now\n3. Proposed Solution\n4. Out of Scope\n5. User Stories (as sprint tickets)\n6. Acceptance Criteria\n7. Technical Notes & Dependencies\n8. Success Metrics\n9. Timeline & Milestones\n10. Open Questions""",
        "jtbd": """1. Job Story Overview\n2. Functional Jobs (core tasks)\n3. Emotional Jobs (how users want to feel)\n4. Social Jobs (how users want to be perceived)\n5. Job Map (step-by-step execution)\n6. Opportunity Analysis (underserved outcomes)\n7. Solution Requirements (anchored to jobs)\n8. Success Metrics (outcome-based only)\n9. Constraints & Assumptions\n10. Open Questions""",
        "hypothesis": """1. Problem Hypothesis\n2. Customer Segment & Early Adopters\n3. Value Proposition Hypothesis\n4. Feature Bets (each as a testable hypothesis)\n5. Riskiest Assumptions\n6. Experiment Design & Test Plan\n7. Minimum Viable Scope\n8. Success & Failure Criteria\n9. Learning Plan & Pivot Triggers\n10. Open Questions""",
        "rfc": """1. Summary & Motivation\n2. Background & Prior Art\n3. Proposed Design (with alternatives considered)\n4. System Architecture & Data Models\n5. API Contracts & Interface Definitions\n6. Performance, Scalability & SLOs\n7. Security & Privacy Considerations\n8. Migration, Rollback & Operational Plan\n9. Open Questions & Decisions Needed\n10. References""",
        "apple": """1. Vision & Experience Narrative\n2. Customer Problem (told as a story)\n3. Design Principles & Hard Constraints\n4. Core Experience Definition\n5. Feature Scope (what's in — and explicitly out)\n6. Privacy, Security & Ecosystem Fit\n7. Platform & Ecosystem Integration\n8. Quality Bar & Performance Requirements\n9. Launch Readiness Criteria\n10. Open Questions""",
        "microsoft": """1. Executive Summary & Business Justification\n2. Customer & IT Admin Requirements\n3. Compliance, Security & Accessibility (WCAG, SOC2, ISO 27001)\n4. Feature Requirements\n5. Microsoft 365 / Azure / Teams Integration\n6. Technical Architecture & Compatibility\n7. Deployment, Admin Controls & Change Management\n8. Success Metrics (end-user NPS + IT admin adoption)\n9. Risks & Dependencies\n10. Open Questions""",
        "meta": """1. Executive Summary\n2. Social Graph & Network Effect Opportunity\n3. User Personas & Behavioral Signals\n4. Feature Requirements\n5. A/B Testing & Experiment Design Plan\n6. Engagement & Viral Mechanics\n7. Algorithmic Distribution Implications\n8. Success Metrics (DAU, MAU, Engagement Rate, k-factor)\n9. Risks & Mitigations\n10. Open Questions""",
        "nvidia": """1. Executive Summary\n2. Performance Specification (TFLOPS, latency, bandwidth, power)\n3. Developer & Ecosystem Requirements\n4. Feature Requirements\n5. CUDA / SDK / Toolchain Integration\n6. Hardware-Software Co-Design Constraints\n7. Benchmark & Competitive Analysis\n8. Developer Experience & Documentation Requirements\n9. Risks & Dependencies\n10. Open Questions""",
        "openai": """1. Executive Summary\n2. Capability Requirements\n3. Safety Requirements (equal weight to capability)\n4. Red-Teaming & Misuse Vector Analysis\n5. Evaluation Benchmarks & Model Criteria\n6. Staged Rollout & Deployment Gates\n7. Responsible Use Policy Requirements\n8. Success Metrics (capability × safety)\n9. Risks & Open Questions\n10. What 'Ready to Ship' Means""",
        "anthropic": """1. Executive Summary\n2. Product Capability Requirements\n3. Safety Case (harmlessness assessment)\n4. Honesty & Transparency Requirements\n5. Helpfulness Measures\n6. Constitutional AI Alignment Criteria\n7. Responsible Scaling Policy Implications\n8. Evaluation Criteria (capability + alignment)\n9. Epistemic Humility — What We Don't Yet Know\n10. Open Questions""",
        "tesla": """1. Requirements Deletion Log (what was challenged and cut)\n2. First-Principles Problem Framing\n3. Product Specification (physics & economics justified)\n4. Software-Defined Feature Requirements\n5. OTA Update & Iteration Strategy\n6. Vertical Integration Requirements\n7. COGS Target & Cost Engineering\n8. Automation & Scale Plan\n9. Performance Benchmarks & Cycle Time Targets\n10. Open Questions""",
        "mercedes": """1. Executive Summary\n2. Luxury Experience & Brand Requirements\n3. Safety Standards (ISO 26262, ASIL ratings)\n4. Feature Specifications\n5. Quality Validation Gates & Homologation\n6. Regulatory & Compliance Requirements\n7. Technical Architecture & Durability\n8. Long-Term Reliability KPIs\n9. Premium UX & Craftsmanship Standards\n10. Open Questions""",
        "honda": """1. Executive Summary (Honda Way framing)\n2. Kaizen Baseline — Current State & Improvement Targets\n3. Customer Requirements (reliability-first)\n4. Feature Requirements\n5. Manufacturing Feasibility & Total Cost of Ownership\n6. Quality Standards (defect rate targets, serviceability)\n7. Global Regulatory Compliance\n8. Supplier & Supply Chain Requirements\n9. Risks & Continuous Improvement Plan\n10. Open Questions""",
        "elililly": """1. Unmet Medical Need & Clinical Rationale\n2. Patient & Healthcare Provider Requirements\n3. Regulatory Strategy (FDA 21 CFR, ICH E6)\n4. Clinical Evidence Requirements (endpoints, safety)\n5. Pharmacovigilance Obligations\n6. Feature / Product Specifications\n7. Health Economics & Market Access (QALY, cost-effectiveness)\n8. Patient Outcome Metrics (north star)\n9. Risks & Regulatory Contingencies\n10. Open Questions""",
        "novartis": """1. Unmet Medical Need\n2. Clinical Development Rationale\n3. Patient-Centricity Requirements (PRO instruments)\n4. Regulatory Pathway (FDA, EMA, PMDA)\n5. Product / Feature Specifications\n6. Medical Affairs & Real-World Evidence Strategy\n7. Pharmacovigilance & Safety Obligations\n8. Market Access Criteria\n9. Success Metrics (clinical outcomes + quality of life)\n10. Open Questions""",
        "exxon": """1. Executive Summary\n2. HSE Impact Assessment (Health, Safety, Environment)\n3. Process Safety Requirements (OSHA PSM 1910.119)\n4. Operational Requirements\n5. Regulatory Compliance (EPA, FERC)\n6. CAPEX / OPEX Analysis\n7. Upstream / Downstream Integration\n8. ESG Reporting & Emissions Requirements\n9. Success Metrics (safety KPIs, efficiency, asset value)\n10. Open Questions""",
        "chevron": """1. Executive Summary\n2. Safety & PSM Standards (non-negotiable)\n3. Environmental Stewardship Requirements\n4. Feature / Operational Requirements\n5. Energy Transition & Net-Zero Alignment\n6. CAPEX / OPEX & Supply Chain Analysis\n7. Regulatory Compliance\n8. ESG Impact Measurement Framework\n9. Success Metrics (TRIR, environmental KPIs, financial returns)\n10. Open Questions""",
        "maersk": """1. Executive Summary\n2. Supply Chain Resilience Requirements\n3. End-to-End Visibility & Tracking Requirements\n4. Multi-Modal Integration (ocean, inland, air)\n5. Feature Specifications\n6. Decarbonisation Requirements (net-zero 2040)\n7. Trade Lane Economics & Operational KPIs\n8. Digital-Physical Integration Requirements\n9. Sustainability Metrics (CO2 per TEU-km)\n10. Open Questions""",
        "jpmorgan": """1. Executive Summary\n2. Regulatory Compliance (Basel III, Dodd-Frank, MiFID II, GDPR)\n3. Risk Assessment & Financial Controls Design\n4. Feature Requirements\n5. Audit Trail & Model Risk Management\n6. Security & Fraud Prevention Architecture\n7. Operational Resilience Requirements\n8. Success Metrics (risk-adjusted returns, compliance adherence)\n9. Dependencies & Integration Requirements\n10. Open Questions""",
        "visa": """1. Executive Summary\n2. Availability & Latency SLOs (99.999%, sub-100ms)\n3. Fraud Detection & Prevention Requirements\n4. Feature Specifications\n5. PCI-DSS, PSD2 & Tokenisation Architecture\n6. Global Regulatory Variation (EMV, local payment schemes)\n7. Authorization Rate & Acceptance Optimization\n8. Security Architecture Requirements\n9. Success Metrics (auth rate, fraud basis points, acceptance)\n10. Open Questions""",
        "spacex": """1. First-Principles Problem Framing\n2. Requirements Deletion & Simplification Log\n3. Physics & Economics Justification\n4. Product Specifications (cost + performance targets)\n5. Reusability & Rapid Iteration Requirements\n6. Reliability & Safety (FMEA summary)\n7. Cost-Per-Unit Targets\n8. Manufacturing & Cycle Time Goals\n9. Success Metrics (cost, reliability, iteration speed)\n10. Open Questions""",
        "boeing": """1. Executive Summary\n2. Airworthiness & Safety Requirements (FAA/EASA)\n3. DO-178C Software Criticality Levels (DAL A–E)\n4. System Redundancy Architecture\n5. Feature & Functional Requirements\n6. FMEA Summary (Failure Mode & Effects Analysis)\n7. Design Assurance Requirements\n8. Certification Timeline & Regulatory Milestones\n9. Traceability Matrix (requirement to test case)\n10. Open Questions""",
    }
    sections_list = format_sections_map.get(format_style, format_sections_map["google"])

    prompt = f"""You are an expert product manager writing a Product Requirements Document (PRD).

{format_instructions}

Product Details:
- Product Name: {product_name}
- Problem Statement: {problem}
- Target Users: {target_users or 'Not specified'}
- Key Features: {key_features or 'Not specified'}
- Success Metrics: {success_metrics or 'Not specified'}
- Company Stage: {company_stage or 'Not specified'}
- Additional Context: {additional_context or 'None'}

STRICT RULES — non-negotiable:
1. Start with a Table of Contents listing all 10 sections.
2. Write ALL 10 sections below. Every section is mandatory.
3. Each section must be a MAXIMUM of {words_per_section} words. Never cut a section — cut detail instead.
4. Depth scales with the limit: concise for Brief, balanced for Medium, detailed for Extensive.
5. The document MUST end with a completed "10. Open Questions" section.

Sections (each max {words_per_section} words):
{sections_list}

Write in Markdown. Be professional, specific, and actionable."""

    user_id = user["id"]

    # ── Streaming SSE generator ────────────────────────────────────────────────
    async def event_stream():
        full_content = ""
        try:
            _client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

            # Round 1 — stream main generation
            async with _client.messages.stream(
                model="claude-sonnet-4-6", max_tokens=max_tok,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                async for text in stream.text_stream:
                    full_content += text
                    yield f"data: {json.dumps({'text': text})}\n\n"
                final_msg = await stream.get_final_message()

            # Round 2 — only if model hit the token wall
            if final_msg.stop_reason == "max_tokens":
                continuation = (
                    "The document was cut off. You have 1500 tokens to finish it. "
                    "Identify every section from the required list that is missing or incomplete. "
                    "Write each remaining section — condense aggressively if needed. "
                    "COMPLETENESS IS MANDATORY. Do NOT repeat anything already written. "
                    "End with a completed Open Questions section."
                )
                async with _client.messages.stream(
                    model="claude-sonnet-4-6", max_tokens=1500,
                    messages=[
                        {"role": "user",      "content": prompt},
                        {"role": "assistant", "content": full_content},
                        {"role": "user",      "content": continuation},
                    ]
                ) as stream2:
                    async for text in stream2.text_stream:
                        full_content += text
                        yield f"data: {json.dumps({'text': text})}\n\n"

        except Exception as e:
            logger.error(f"Anthropic stream error: {e}")
            yield f"data: {json.dumps({'error': 'AI generation failed. Please try again.'})}\n\n"
            return

        # Save to DB and deduct credits after streaming completes
        prd_id = str(uuid.uuid4())
        try:
            conn = get_db()
            conn.execute(
                "INSERT INTO prds (id, user_id, title, content, format_style, target_users, "
                "key_features, success_metrics, company_stage, additional_context) "
                "VALUES (?,?,?,?,?,?,?,?,?,?)",
                (prd_id, user_id, product_name, full_content, format_style,
                 target_users, key_features, success_metrics, company_stage, additional_context)
            )
            conn.execute(
                "UPDATE users SET credits_used = credits_used + ?, "
                "prds_used_this_month = prds_used_this_month + 1 WHERE id=?",
                (credit_cost, user_id)
            )
            conn.commit(); conn.close()
        except Exception as e:
            logger.error(f"DB save error after stream: {e}")

        new_remaining = max(0, credits_remaining - credit_cost)
        yield f"data: {json.dumps({'done': True, 'prd_id': prd_id, 'title': product_name, 'credits_remaining': new_remaining})}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"},
    )

@app.get("/api/prd/{prd_id}")
async def get_prd(prd_id: str, request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    conn = get_db()
    prd  = conn.execute("SELECT * FROM prds WHERE id=? AND user_id=?", (prd_id, user["id"])).fetchone()
    conn.close()
    if not prd: raise HTTPException(status_code=404, detail="PRD not found")
    return JSONResponse(dict(prd))

@app.get("/api/prd/{prd_id}/download")
async def download_prd(prd_id: str, request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    conn = get_db()
    prd  = conn.execute("SELECT * FROM prds WHERE id=? AND user_id=?", (prd_id, user["id"])).fetchone()
    conn.close()
    if not prd: raise HTTPException(status_code=404, detail="PRD not found")
    docx_bytes = markdown_to_docx(prd["content"])
    filename   = re.sub(r'[^\w\s-]', '', prd["title"]).strip().replace(' ', '_')
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}_PRD.docx"'}
    )

@app.delete("/api/prd/{prd_id}")
async def delete_prd(prd_id: str, request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    conn = get_db()
    conn.execute("DELETE FROM prds WHERE id=? AND user_id=?", (prd_id, user["id"]))
    conn.commit(); conn.close()
    return JSONResponse({"success": True})

# ─── Stripe API ───────────────────────────────────────────────────────────────
@app.post("/api/create-checkout-session")
async def create_checkout_session(request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{"price": STRIPE_PRICE_ID, "quantity": 1}],
            mode="subscription",
            customer_email=user["email"],
            success_url=f"{BASE_URL}/upgrade/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{BASE_URL}/upgrade",
            metadata={"user_id": user["id"]}
        )
        return JSONResponse({"checkout_url": session.url})
    except Exception as e:
        logger.error(f"Stripe error: {e}")
        return JSONResponse({"error": "Payment setup failed."}, status_code=500)


@app.post("/api/create-yearly-checkout-session")
async def create_yearly_checkout_session(request: Request):
    user = get_current_user(request)
    if not user: return JSONResponse({"error": "Not authenticated"}, status_code=401)
    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{"price": STRIPE_YEARLY_PRICE_ID, "quantity": 1}],
            mode="subscription", customer_email=user["email"],
            success_url=f"{BASE_URL}/upgrade/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{BASE_URL}/upgrade",
            metadata={"user_id": user["id"], "plan": "yearly"}
        )
        return JSONResponse({"checkout_url": session.url})
    except Exception as e:
        logger.error(f"Stripe yearly error: {e}")
        return JSONResponse({"error": "Payment setup failed."}, status_code=500)
@app.post("/api/stripe-webhook")
async def stripe_webhook(request: Request):
    payload    = await request.body()
    sig_header = request.headers.get("stripe-signature", "")
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    if event["type"] == "checkout.session.completed":
        s       = event["data"]["object"]
        user_id = s.get("metadata", {}).get("user_id")
        if user_id:
            conn = get_db()
            conn.execute(
                "UPDATE users SET plan=?, stripe_customer_id=?, stripe_subscription_id=? WHERE id=?",
                (s.get("metadata",{}).get("plan","pro"), s.get("customer"), s.get("subscription"), user_id)
            )
            conn.commit(); conn.close()
    elif event["type"] in ("customer.subscription.deleted", "customer.subscription.paused"):
        sub  = event["data"]["object"]
        conn = get_db()
        conn.execute("UPDATE users SET plan='free' WHERE stripe_subscription_id=?", (sub["id"],))
        conn.commit(); conn.close()
    return JSONResponse({"received": True})

# ─── Admin API ────────────────────────────────────────────────────────────────
@app.get("/api/admin/stats")
async def admin_stats(request: Request):
    user = get_current_user(request)
    if not user or user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Forbidden")
    conn         = get_db()
    total_users  = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    pro_users    = conn.execute("SELECT COUNT(*) FROM users WHERE plan='pro'").fetchone()[0]
    total_prds   = conn.execute("SELECT COUNT(*) FROM prds").fetchone()[0]
    recent_users = [dict(r) for r in conn.execute(
        "SELECT id, name, email, plan, role, prds_used_this_month, created_at "
        "FROM users ORDER BY created_at DESC LIMIT 20"
    ).fetchall()]
    recent_prds  = [dict(r) for r in conn.execute(
        "SELECT p.id, p.title, p.created_at, u.email "
        "FROM prds p JOIN users u ON p.user_id=u.id ORDER BY p.created_at DESC LIMIT 20"
    ).fetchall()]
    conn.close()
    return JSONResponse({
        "total_users": total_users, "pro_users": pro_users,
        "total_prds": total_prds,   "mrr": pro_users * 20,
        "users": recent_users,      "prds": recent_prds
    })
